from typing import List, Dict, Any, Optional, Tuple
import os
import re
import logging
from datetime import datetime
import traceback

import pandas as pd
import pdfplumber
from PIL import Image
import pytesseract
from docx import Document
import io

logger = logging.getLogger(__name__)

# Настройка OCR (опционально). Оставьте как есть, если путь отличается в вашей системе.
try:
    # Попробуем несколько возможных путей к Tesseract
    possible_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract"
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            logger.info(f"Tesseract найден по пути: {path}")
            break
    else:
        logger.warning("Tesseract не найден. OCR для изображений будет недоступен.")
except Exception as e:
    logger.error(f"Ошибка настройки Tesseract: {e}")


def parse_tariff_file(file_path: str, supplier_id: int) -> List[Dict[str, Any]]:
    """
    Парсинг файлов тарифов с автоматическим определением типа транспорта.
    Поддерживает: XLS, XLSX, CSV, DOCX, PDF (текст), PDF (скан), JPG, PNG.
    """
    logger.info(f"Начинаем парсинг файла: {file_path}")
    
    try:
        # Используем фабрику парсеров для автоматического определения типа транспорта
        from parser_factory import ParserFactory
        
        # Определяем тип транспорта и парсим данные
        result = ParserFactory.parse_with_auto_detection(file_path, supplier_id)
        
        logger.info(f"Парсинг завершен. Извлечено {len(result)} записей")
        return result
        
    except Exception as e:
        logger.error(f"Ошибка парсинга файла {file_path}: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return []


def extract_text_from_file(file_path: str) -> str:
    """
    Универсальная функция для извлечения текста из файлов различных форматов.
    Возвращает извлеченный текст в виде строки.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext in [".xls", ".xlsx", ".csv"]:
            return _extract_text_from_excel(file_path)
        elif ext == ".pdf":
            return _extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return _extract_text_from_docx(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            return _extract_text_from_image(file_path)
        else:
            logger.warning(f"Неподдерживаемый формат файла для извлечения текста: {ext}")
            return ""
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из файла {file_path}: {e}")
        return ""


def _extract_text_from_excel(file_path: str) -> str:
    """Извлечение текста из Excel/CSV файлов."""
    try:
        logger.info(f"Извлекаем текст из Excel файла: {file_path}")
        
        # Пробуем разные движки для чтения Excel
        df = None
        engines = ["openpyxl", "xlrd"]
        
        for engine in engines:
            try:
                df = pd.read_excel(file_path, engine=engine)
                logger.info(f"Успешно прочитан файл с движком {engine}")
                break
            except Exception as e:
                logger.warning(f"Ошибка с движком {engine}: {e}")
                continue
        
        # Если Excel не удался, пробуем CSV
        if df is None:
            try:
                df = pd.read_csv(file_path)
                logger.info("Успешно прочитан как CSV")
            except Exception as e:
                logger.error(f"Ошибка чтения CSV: {e}")
                return ""
        
        if df is None or df.empty:
            return ""
        
        # Конвертируем DataFrame в текст
        text_lines = []
        
        # Добавляем заголовки
        headers = [str(col) for col in df.columns]
        text_lines.append(" | ".join(headers))
        
        # Добавляем данные
        for _, row in df.iterrows():
            row_text = " | ".join([str(val) if pd.notna(val) else "" for val in row])
            text_lines.append(row_text)
        
        return "\n".join(text_lines)
        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из Excel: {e}")
        return ""


def _extract_text_from_pdf(file_path: str) -> str:
    """Извлечение текста из PDF файлов."""
    text_lines = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                logger.info(f"Обрабатываем страницу {page_num + 1}")
                
                # Пробуем извлечь текст
                text = page.extract_text()
                
                if text:
                    text_lines.append(f"=== Страница {page_num + 1} ===")
                    text_lines.append(text)
                else:
                    # OCR для сканов
                    try:
                        im = page.to_image()
                        pil_image = Image.fromarray(im.original)
                        ocr_text = pytesseract.image_to_string(pil_image, lang="rus+eng")
                        if ocr_text:
                            text_lines.append(f"=== Страница {page_num + 1} (OCR) ===")
                            text_lines.append(ocr_text)
                    except Exception as e:
                        logger.warning(f"Ошибка OCR для страницы {page_num + 1}: {e}")
                        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из PDF: {e}")
    
    return "\n".join(text_lines)


def _extract_text_from_docx(file_path: str) -> str:
    """Извлечение текста из DOCX файлов."""
    text_lines = []
    
    try:
        logger.info(f"Извлекаем текст из DOCX файла: {file_path}")
        
        doc = Document(file_path)
        
        # Извлекаем текст из параграфов
        for para in doc.paragraphs:
            if para.text.strip():
                text_lines.append(para.text)
        
        # Извлекаем текст из таблиц
        for table_idx, table in enumerate(doc.tables):
            text_lines.append(f"=== Таблица {table_idx + 1} ===")
            
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_lines.append(row_text)
        
        return "\n".join(text_lines)
        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из DOCX: {e}")
        return ""


def _extract_text_from_image(file_path: str) -> str:
    """Извлечение текста из изображений с помощью OCR."""
    try:
        logger.info(f"Извлекаем текст из изображения: {file_path}")
        
        # Проверяем, доступен ли Tesseract
        try:
            pytesseract.get_tesseract_version()
        except Exception as e:
            logger.error(f"Tesseract недоступен: {e}")
            return ""
        
        image = Image.open(file_path)
        
        # Пробуем разные настройки OCR для лучшего распознавания
        configs = [
            '--oem 3 --psm 6',  # Стандартные настройки
            '--oem 3 --psm 3',  # Автоматическое определение страницы
            '--oem 1 --psm 6',  # Legacy OCR engine
            '--oem 3 --psm 8',  # Одна строка текста
            '--oem 3 --psm 13', # Необработанная строка
        ]
        
        best_text = ""
        for config in configs:
            try:
                text = pytesseract.image_to_string(image, lang="rus+eng", config=config)
                if text and len(text.strip()) > len(best_text.strip()):
                    best_text = text
                    logger.info(f"Успешное распознавание с конфигурацией: {config}")
            except Exception as e:
                logger.warning(f"Ошибка OCR с конфигурацией {config}: {e}")
        
        return best_text
        
    except Exception as e:
        logger.error(f"Ошибка OCR для изображения {file_path}: {e}")
        return ""


def _parse_excel(file_path: str) -> List[Dict[str, Any]]:
    """Парсинг Excel/CSV с нормализацией заголовков."""
    try:
        logger.info(f"Начинаем парсинг Excel файла: {file_path}")
        
        # Пробуем разные движки для чтения Excel
        df = None
        engines = ["openpyxl", "xlrd"]
        
        for engine in engines:
            try:
                logger.info(f"Пробуем движок: {engine}")
                df = pd.read_excel(file_path, engine=engine)
                logger.info(f"Успешно прочитан файл с движком {engine}")
                break
            except Exception as e:
                logger.warning(f"Ошибка с движком {engine}: {e}")
                continue
        
        # Если Excel не удался, пробуем CSV
        if df is None:
            try:
                logger.info("Пробуем прочитать как CSV")
                df = pd.read_csv(file_path)
                logger.info("Успешно прочитан как CSV")
            except Exception as e:
                logger.error(f"Ошибка чтения CSV: {e}")
                return []
        
        if df is None or df.empty:
            logger.warning("Файл пустой или не содержит данных")
            return []
        
        logger.info(f"Размер данных: {df.shape[0]} строк, {df.shape[1]} колонок")
        logger.info(f"Колонки: {list(df.columns)}")
        
        # Нормализуем заголовки
        df.columns = [str(c).strip().lower() for c in df.columns]
        logger.info(f"Нормализованные колонки: {list(df.columns)}")
        
        rows: List[Dict[str, Any]] = []
        for idx, r in df.iterrows():
            logger.debug(f"Обрабатываем строку {idx + 1}: {r.to_dict()}")
            row = _normalize_row_from_dict(r.to_dict())
            if row:
                rows.append(row)
                logger.debug(f"Добавлена строка: {row}")
            else:
                logger.debug(f"Строка {idx + 1} пропущена (пустая)")
        
        logger.info(f"Извлечено {len(rows)} записей из Excel файла")
        return rows
        
    except Exception as e:
        logger.error(f"Ошибка чтения Excel/CSV: {e}")
        return []


def _parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Парсинг PDF (таблицы + OCR fallback)."""
    rows: List[Dict[str, Any]] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = None
                try:
                    text = page.extract_text()
                except Exception:
                    text = None

                if text:
                    # Пробуем таблицу
                    try:
                        table = page.extract_table()
                    except Exception:
                        table = None

                    if table and len(table) > 1:
                        headers = [str(h).strip().lower() for h in table[0]]
                        for line in table[1:]:
                            r = {headers[i]: line[i] for i in range(min(len(headers), len(line)))}
                            row = _normalize_row_from_dict(r)
                            if row:
                                rows.append(row)
                    else:
                        rows.extend(_extract_freestyle_text(text))
                else:
                    # OCR для сканов
                    try:
                        im = page.to_image()
                        pil_image = Image.fromarray(im.original)
                        text = pytesseract.image_to_string(pil_image, lang="rus+eng")
                        if text:
                            rows.extend(_extract_freestyle_text(text))
                    except Exception:
                        pass
    except Exception as e:
        logger.error(f"Ошибка парсинга PDF: {e}")
    return rows


def _parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """Парсинг DOCX (таблицы + свободный текст)."""
    rows: List[Dict[str, Any]] = []
    try:
        logger.info(f"Начинаем парсинг DOCX файла: {file_path}")
        
        doc = Document(file_path)
        logger.info(f"Документ загружен, количество таблиц: {len(doc.tables)}")
        
        # Парсим таблицы
        for table_idx, tbl in enumerate(doc.tables):
            logger.info(f"Обрабатываем таблицу {table_idx + 1}")
            
            if not tbl.rows:
                logger.warning(f"Таблица {table_idx + 1} пустая")
                continue
                
            # Извлекаем заголовки
            headers = [cell.text.strip().lower() for cell in tbl.rows[0].cells]
            logger.info(f"Заголовки таблицы: {headers}")
            
            # Обрабатываем строки данных
            for row_idx, row in enumerate(tbl.rows[1:], 1):
                data = [cell.text.strip() for cell in row.cells]
                logger.debug(f"Строка {row_idx}: {data}")
                
                # Создаем словарь данных
                d = {headers[i]: data[i] for i in range(min(len(headers), len(data)))}
                row_data = _normalize_row_from_dict(d)
                if row_data:
                    rows.append(row_data)
                    logger.debug(f"Добавлена строка из таблицы: {row_data}")
                else:
                    logger.debug(f"Строка {row_idx} из таблицы пропущена")

        # Парсим свободный текст
        logger.info("Обрабатываем свободный текст")
        text = "\n".join(p.text for p in doc.paragraphs)
        logger.info(f"Длина текста: {len(text)} символов")
        
        if text.strip():
            freestyle_rows = _extract_freestyle_text(text)
            rows.extend(freestyle_rows)
            logger.info(f"Извлечено {len(freestyle_rows)} записей из свободного текста")
        
        logger.info(f"Всего извлечено {len(rows)} записей из DOCX файла")
        
    except Exception as e:
        logger.error(f"Ошибка парсинга DOCX: {e}")
    return rows


def _parse_image(file_path: str) -> List[Dict[str, Any]]:
    """Парсинг изображений с помощью OCR."""
    rows: List[Dict[str, Any]] = []
    try:
        logger.info(f"Начинаем парсинг изображения: {file_path}")
        
        # Проверяем, доступен ли Tesseract
        try:
            pytesseract.get_tesseract_version()
        except Exception as e:
            logger.error(f"Tesseract недоступен: {e}")
            return rows
        
        image = Image.open(file_path)
        logger.info(f"Изображение загружено: {image.size} пикселей")
        
        # Пробуем разные настройки OCR для лучшего распознавания
        configs = [
            '--oem 3 --psm 6',  # Стандартные настройки
            '--oem 3 --psm 3',  # Автоматическое определение страницы
            '--oem 1 --psm 6',  # Legacy OCR engine
        ]
        
        best_text = ""
        for config in configs:
            try:
                text = pytesseract.image_to_string(image, lang="rus+eng", config=config)
                if text and len(text.strip()) > len(best_text.strip()):
                    best_text = text
                    logger.info(f"Успешное распознавание с конфигурацией: {config}")
            except Exception as e:
                logger.warning(f"Ошибка OCR с конфигурацией {config}: {e}")
        
        if best_text:
            logger.info(f"Распознанный текст (первые 200 символов): {best_text[:200]}...")
            extracted_rows = _extract_freestyle_text(best_text)
            rows.extend(extracted_rows)
            logger.info(f"Извлечено {len(extracted_rows)} записей из изображения")
        else:
            logger.warning("Не удалось распознать текст из изображения")
            
    except Exception as e:
        logger.error(f"Ошибка OCR для изображения {file_path}: {e}")
    
    return rows


def _normalize_row_from_dict(d: Dict[str, Any]) -> Dict[str, Any]:
    """Приведение данных строки к единому формату."""
    if not isinstance(d, dict):
        logger.debug("Не словарь, пропускаем")
        return {}
    
    # Проверяем, есть ли хоть какие-то данные
    has_data = any(str(v).strip() for v in d.values() if v is not None and str(v).strip())
    if not has_data:
        logger.debug("Строка пустая, пропускаем")
        return {}
    
    logger.debug(f"Нормализуем данные: {d}")
    
    # Ищем данные в разных возможных полях
    transport_type = _map_transport(
        d.get("transport") or d.get("mode") or d.get("тип") or 
        d.get("transport_type") or d.get("вид_транспорта") or "auto"
    )
    
    basis = str(d.get("basis") or d.get("incoterm") or d.get("базис") or 
                d.get("incoterms") or "EXW").upper()
    
    origin_country = (d.get("origin_country") or d.get("страна_отправления") or 
                     d.get("country_from") or d.get("from_country") or "")
    
    origin_city = (d.get("origin_city") or d.get("from") or d.get("origin") or 
                  d.get("город_отправления") or d.get("city_from") or d.get("from_city") or 
                  d.get("origin_city") or "")
    
    destination_country = (d.get("destination_country") or d.get("страна_назначения") or 
                          d.get("country_to") or d.get("to_country") or "")
    
    destination_city = (d.get("destination_city") or d.get("to") or d.get("destination") or 
                       d.get("город_назначения") or d.get("city_to") or d.get("to_city") or "")
    
    vehicle_type = (d.get("vehicle_type") or d.get("container") or d.get("тип_тс") or 
                   d.get("vehicle") or d.get("транспорт") or "")
    
    price_rub = _float_or_none(d.get("price_rub") or d.get("rub") or d.get("руб") or 
                              d.get("price") or d.get("стоимость"))
    
    price_usd = _float_or_none(d.get("price_usd") or d.get("usd") or d.get("доллар") or 
                              d.get("price_usd") or d.get("стоимость_usd"))
    
    validity_date = _date_or_none(d.get("validity") or d.get("valid_to") or d.get("date") or 
                                 d.get("дата") or d.get("срок_действия"))
    
    transit_time_days = _int_or_none(d.get("transit_days") or d.get("tt") or d.get("срок") or 
                                   d.get("время_в_пути") or d.get("transit_time"))
    
    result = {
        "transport_type": transport_type,
        "basis": basis,
        "origin_country": origin_country,
        "origin_city": origin_city,
        "destination_country": destination_country,
        "destination_city": destination_city,
        "vehicle_type": vehicle_type,
        "price_rub": price_rub,
        "price_usd": price_usd,
        "validity_date": validity_date,
        "transit_time_days": transit_time_days,
    }
    
    # Убираем пустые значения
    result = {k: v for k, v in result.items() if v is not None and v != ""}
    
    logger.debug(f"Результат нормализации: {result}")
    return result


def _extract_freestyle_text(text: str) -> List[Dict[str, Any]]:
    """Улучшенное извлечение из свободного текста (эвристики)."""
    rows: List[Dict[str, Any]] = []
    if not text:
        return rows

    logger.info(f"Анализируем текст длиной {len(text)} символов")
    
    # Разбиваем на строки и анализируем каждую
    lines = text.splitlines()
    logger.info(f"Найдено {len(lines)} строк текста")
    
    for i, line in enumerate(lines):
        l = line.strip()
        if not l or len(l) < 5:  # Пропускаем слишком короткие строки
            continue

        logger.debug(f"Анализируем строку {i+1}: {l[:100]}...")
        
        # Ищем цены в разных форматах
        price_patterns = [
            # USD цены
            r"USD\s*(\d+(?:[.,]\d+)?)",  # USD 1000
            r"(\d+(?:[.,]\d+)?)\s*USD",  # 1000 USD
            r"(\d+(?:[.,]\d+)?)\s*\$",   # 1000 $
            r"\$\s*(\d+(?:[.,]\d+)?)",   # $ 1000
            r"(\d+(?:[.,]\d+)?)\s*доллар", # 1000 доллар
            r"price.*?(\d+(?:[.,]\d+)?)", # price 1000
            r"rate.*?(\d+(?:[.,]\d+)?)",  # rate 1000
            r"(\d+(?:[.,]\d+)?)\s*/\s*(?:40|20|container|bl)", # 1000/40HC
            
            # Рублевые цены
            r"(\d+(?:[.,]\d+)?)\s*руб",   # 1000 руб
            r"(\d+(?:[.,]\d+)?)\s*₽",     # 1000 ₽
            r"(\d+(?:[.,]\d+)?)\s*RUB",   # 1000 RUB
            
            # RMB цены
            r"(\d+(?:[.,]\d+)?)\s*RMB",   # 1000 RMB
            r"RMB\s*(\d+(?:[.,]\d+)?)",   # RMB 1000
        ]
        
        price_usd = None
        price_rub = None
        
        for pattern in price_patterns:
            match = re.search(pattern, l, re.IGNORECASE)
            if match:
                try:
                    price_str = match.group(1).replace(',', '.').replace(' ', '')
                    price = float(price_str)
                    
                    # Фильтруем слишком маленькие цены (менее 1)
                    if price < 1:
                        continue
                        
                    if 'руб' in pattern or '₽' in pattern or 'RUB' in pattern:
                        price_rub = price
                    elif 'RMB' in pattern:
                        # Конвертируем RMB в USD (примерный курс)
                        price_usd = price / 7.2
                    else:
                        price_usd = price
                    logger.debug(f"Найдена цена: {price} (USD: {price_usd}, RUB: {price_rub})")
                    break
                except ValueError:
                    continue
        
        # Ищем маршруты в разных форматах
        route_patterns = [
            # Русские города
            r"([А-Яа-я\s]+)\s*[-–—]\s*([А-Яа-я\s]+)",  # Москва - Санкт-Петербург
            r"([А-Яа-я\s]+)\s*→\s*([А-Яа-я\s]+)",      # Москва → Санкт-Петербург
            r"([А-Яа-я\s]+)\s*>\s*([А-Яа-я\s]+)",      # Москва > Санкт-Петербург
            r"([А-Яа-я\s]+)\s*до\s*([А-Яа-я\s]+)",     # Москва до Санкт-Петербург
            
            # Английские города
            r"([A-Za-z\s]+)\s*[-–—]\s*([A-Za-z\s]+)",  # Beijing - Moscow
            r"([A-Za-z\s]+)\s*→\s*([A-Za-z\s]+)",      # Beijing → Moscow
            r"([A-Za-z\s]+)\s*>\s*([A-Za-z\s]+)",      # Beijing > Moscow
            r"([A-Za-z\s]+)\s*to\s*([A-Za-z\s]+)",     # Beijing to Moscow
            
            # Специфичные паттерны для логистики
            r"from\s+([A-Za-z\s]+)\s+to\s+([A-Za-z\s]+)",  # from Beijing to Moscow
            r"([A-Za-z\s]+)\s*-\s*([A-Za-z\s]+)\s*route",  # Beijing-Moscow route
            r"([A-Za-z\s]+)\s*/\s*([A-Za-z\s]+)",      # Beijing/Moscow
        ]
        
        origin = None
        destination = None
        
        for pattern in route_patterns:
            match = re.search(pattern, l)
            if match:
                origin = match.group(1).strip()
                destination = match.group(2).strip()
                
                # Фильтруем слишком короткие или нерелевантные названия
                if len(origin) > 2 and len(destination) > 2:
                    # Убираем общие слова
                    exclude_words = ['route', 'from', 'to', 'до', 'от', 'маршрут', 'путь']
                    origin_clean = ' '.join([word for word in origin.split() if word.lower() not in exclude_words])
                    destination_clean = ' '.join([word for word in destination.split() if word.lower() not in exclude_words])
                    
                    if origin_clean and destination_clean:
                        origin = origin_clean
                        destination = destination_clean
                        logger.debug(f"Найден маршрут: {origin} -> {destination}")
                        break
        
        # Ищем тип транспорта
        transport_type = "auto"  # по умолчанию
        l_lower = l.lower()
        
        # Железнодорожный транспорт
        if any(word in l_lower for word in ['жд', 'железнодорожный', 'rail', 'fob.rail', 'lcl.rail', 'station', 'coc']):
            transport_type = "rail"
        # Морской транспорт
        elif any(word in l_lower for word in ['море', 'морской', 'sea', 'fcl', 'pol', 'pop', 'port', 'carrier', 'feeder']):
            transport_type = "sea"
        # Авиатранспорт
        elif any(word in l_lower for word in ['авиа', 'воздушный', 'air', 'peking', 'pek', 'svo', 'quotation']):
            transport_type = "air"
        # Мультимодальный
        elif any(word in l_lower for word in ['мульти', 'multimodal', 'mmp']):
            transport_type = "multimodal"
        # Автомобильный (по умолчанию)
        elif any(word in l_lower for word in ['авто', 'автомобильный', 'truck', 'ftl', 'ltl', 'trucking', 'door-to-door']):
            transport_type = "auto"
        
        # Если нашли хотя бы цену или маршрут, создаем запись
        if price_usd or price_rub or (origin and destination):
            row = {
                "transport_type": transport_type,
                "basis": "EXW",
                "origin_city": origin or "Не указан",
                "destination_city": destination or "Не указан",
                "price_usd": price_usd,
                "price_rub": price_rub,
                "vehicle_type": "20т",
            }
            
            # Убираем None значения
            row = {k: v for k, v in row.items() if v is not None}
            
            if row:
                rows.append(row)
                logger.info(f"Создана запись: {row}")
    
    logger.info(f"Всего извлечено записей: {len(rows)}")
    return rows


def _map_transport(transport: str) -> str:
    if not transport:
        return "auto"
    t = str(transport).lower().strip()
    mapping = {
        # Автомобильный транспорт
        "авто": "auto", "автомобильный": "auto", "автомобиль": "auto", "truck": "auto", "автомобильная": "auto",
        "ftl": "auto", "ltl": "auto", "trucking": "auto", "door-to-door": "auto",
        
        # Железнодорожный транспорт
        "жд": "rail", "железнодорожный": "rail", "железная дорога": "rail", "rail": "rail", "railway": "rail", 
        "железнодорожная": "rail", "fob.rail": "rail", "lcl.rail": "rail", "station": "rail", "coc": "rail",
        
        # Морской транспорт
        "море": "sea", "морской": "sea", "морская": "sea", "sea": "sea", "ocean": "sea",
        "fcl": "sea", "pol": "sea", "pop": "sea", "port": "sea", "carrier": "sea", "feeder": "sea",
        
        # Мультимодальный транспорт
        "мульти": "multimodal", "мультимодальный": "multimodal", "мультимодальная": "multimodal", 
        "multimodal": "multimodal", "mmp": "multimodal",
        
        # Авиатранспорт
        "авиа": "air", "авиационный": "air", "авиационная": "air", "air": "air", "airfreight": "air",
        "peking": "air", "pek": "air", "svo": "air", "quotation": "air",
    }
    
    # Проверяем точное совпадение
    if t in mapping:
        return mapping[t]
    
    # Проверяем частичное совпадение
    for key, value in mapping.items():
        if key in t or t in key:
            return value
    
    return "auto"


def _float_or_none(v) -> float:
    if v is None:
        return None
    try:
        return float(str(v).replace(" ", "").replace(",", "."))
    except Exception:
        return None


def _int_or_none(v) -> int:
    if v is None:
        return None
    try:
        return int(float(str(v).replace(" ", "").replace(",", ".")))
    except Exception:
        return None


def _date_or_none(v) -> str:
    if v is None:
        return None
    try:
        if isinstance(v, str):
            s = v.strip()
            if not s:
                return None
            for fmt in ("%d.%m.%Y", "%d/%m/%Y", "%Y-%m-%d", "%d.%m.%y", "%d/%m/%y", "%y-%m-%d"):
                try:
                    return datetime.strptime(s, fmt).strftime("%Y-%m-%d")
                except ValueError:
                    continue
            return s
        # На всякий случай возвращаем строковое представление
        return str(v)
    except Exception:
        return None

