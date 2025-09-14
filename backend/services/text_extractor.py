"""
Сервис для извлечения текста из файлов различных форматов.
Поддерживает: DOCX, XLSX, PDF, PNG, JPEG
"""

import os
import logging
from typing import Dict, Any, Optional, List
import traceback

import pandas as pd
import pdfplumber
from PIL import Image
import pytesseract
from docx import Document

logger = logging.getLogger(__name__)

# Настройка OCR
try:
    possible_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract"
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            logger.info(f"Tesseract найден по пути: {path}")
            break
    else:
        logger.warning("Tesseract не найден. OCR для изображений будет недоступен.")
except Exception as e:
    logger.error(f"Ошибка настройки Tesseract: {e}")


class TextExtractor:
    """Класс для извлечения текста из файлов различных форматов."""
    
    def __init__(self):
        self.supported_formats = {
            '.docx': self._extract_from_docx,
            '.xlsx': self._extract_from_excel,
            '.xls': self._extract_from_excel,
            '.csv': self._extract_from_excel,
            '.pdf': self._extract_from_pdf,
            '.png': self._extract_from_image,
            '.jpg': self._extract_from_image,
            '.jpeg': self._extract_from_image
        }
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Извлекает текст из файла и возвращает структурированный результат.
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Словарь с результатами извлечения:
            {
                'success': bool,
                'text': str,
                'format': str,
                'pages': int,
                'tables': int,
                'images': int,
                'error': str (если есть)
            }
        """
        if not os.path.exists(file_path):
            return {
                'success': False,
                'error': f'Файл не найден: {file_path}',
                'text': '',
                'format': '',
                'pages': 0,
                'tables': 0,
                'images': 0
            }
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.supported_formats:
            return {
                'success': False,
                'error': f'Неподдерживаемый формат файла: {file_ext}',
                'text': '',
                'format': file_ext,
                'pages': 0,
                'tables': 0,
                'images': 0
            }
        
        try:
            logger.info(f"Начинаем извлечение текста из файла: {file_path}")
            result = self.supported_formats[file_ext](file_path)
            result['format'] = file_ext
            result['success'] = True
            logger.info(f"Извлечение завершено успешно: {len(result.get('text', ''))} символов")
            return result
            
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из {file_path}: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'format': file_ext,
                'pages': 0,
                'tables': 0,
                'images': 0
            }
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из DOCX файлов."""
        logger.info(f"Извлекаем текст из DOCX: {file_path}")
        
        doc = Document(file_path)
        text_parts = []
        table_count = 0
        
        # Извлекаем текст из параграфов
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Извлекаем текст из таблиц
        for table_idx, table in enumerate(doc.tables):
            table_count += 1
            text_parts.append(f"\n=== ТАБЛИЦА {table_idx + 1} ===\n")
            
            for row_idx, row in enumerate(table.rows):
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        return {
            'text': '\n'.join(text_parts),
            'pages': 1,  # DOCX обычно одна страница
            'tables': table_count,
            'images': 0
        }
    
    def _extract_from_excel(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из Excel/CSV файлов."""
        logger.info(f"Извлекаем текст из Excel/CSV: {file_path}")
        
        # Пробуем разные движки для чтения Excel
        df = None
        engines = ["openpyxl", "xlrd"]
        
        for engine in engines:
            try:
                # Пробуем разные параметры для чтения
                try:
                    df = pd.read_excel(file_path, engine=engine, header=None)
                except:
                    df = pd.read_excel(file_path, engine=engine, header=0)
                
                logger.info(f"Успешно прочитан с движком {engine}")
                break
            except Exception as e:
                logger.warning(f"Ошибка с движком {engine}: {e}")
                continue
        
        # Если Excel не удался, пробуем CSV с разными кодировками
        if df is None:
            for encoding in ['utf-8', 'cp1251', 'latin1', 'iso-8859-1']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding, header=None)
                    logger.info(f"Успешно прочитан как CSV с кодировкой {encoding}")
                    break
                except Exception as e:
                    logger.warning(f"Ошибка чтения CSV с кодировкой {encoding}: {e}")
                    continue
            
            if df is None:
                raise Exception("Не удалось прочитать файл ни одной кодировкой")
        
        if df is None or df.empty:
            return {
                'text': '',
                'pages': 1,
                'tables': 0,
                'images': 0
            }
        
        # Конвертируем DataFrame в текст с улучшенной обработкой
        text_lines = []
        
        # Добавляем заголовки (если есть)
        if not df.columns.empty:
            headers = []
            for col in df.columns:
                if pd.isna(col) or col == '':
                    headers.append('Unnamed')
                else:
                    headers.append(str(col))
            text_lines.append(" | ".join(headers))
        
        # Добавляем данные с улучшенной обработкой
        for idx, row in df.iterrows():
            row_values = []
            for val in row:
                if pd.isna(val):
                    row_values.append("")
                elif isinstance(val, (int, float)):
                    # Форматируем числа без научной нотации
                    if val == int(val):
                        row_values.append(str(int(val)))
                    else:
                        row_values.append(f"{val:.2f}")
                else:
                    row_values.append(str(val).strip())
            
            row_text = " | ".join(row_values)
            if row_text.strip():  # Добавляем только непустые строки
                text_lines.append(row_text)
        
        return {
            'text': '\n'.join(text_lines),
            'pages': 1,
            'tables': 1,
            'images': 0
        }
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из PDF файлов."""
        logger.info(f"Извлекаем текст из PDF: {file_path}")
        
        text_parts = []
        page_count = 0
        table_count = 0
        image_count = 0
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_count += 1
                    logger.info(f"Обрабатываем страницу {page_num + 1}")
                    
                    # Пробуем извлечь текст
                    text = page.extract_text()
                    
                    if text:
                        text_parts.append(f"\n=== СТРАНИЦА {page_num + 1} ===\n")
                        text_parts.append(text)
                    else:
                        # OCR для сканов
                        try:
                            im = page.to_image()
                            pil_image = Image.fromarray(im.original)
                            ocr_text = pytesseract.image_to_string(pil_image, lang="rus+eng")
                            if ocr_text:
                                text_parts.append(f"\n=== СТРАНИЦА {page_num + 1} (OCR) ===\n")
                                text_parts.append(ocr_text)
                                image_count += 1
                        except Exception as e:
                            logger.warning(f"Ошибка OCR для страницы {page_num + 1}: {e}")
                    
                    # Пробуем извлечь таблицы
                    try:
                        table = page.extract_table()
                        if table and len(table) > 1:
                            table_count += 1
                            text_parts.append(f"\n=== ТАБЛИЦА НА СТРАНИЦЕ {page_num + 1} ===\n")
                            for row in table:
                                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                text_parts.append(row_text)
                    except Exception as e:
                        logger.debug(f"Ошибка извлечения таблицы со страницы {page_num + 1}: {e}")
                        
        except Exception as e:
            logger.error(f"Ошибка обработки PDF: {e}")
            raise
        
        return {
            'text': '\n'.join(text_parts),
            'pages': page_count,
            'tables': table_count,
            'images': image_count
        }
    
    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Извлечение текста из изображений с помощью OCR."""
        logger.info(f"Извлекаем текст из изображения: {file_path}")
        
        # Проверяем, доступен ли Tesseract
        try:
            pytesseract.get_tesseract_version()
        except Exception as e:
            logger.error(f"Tesseract недоступен: {e}")
            raise Exception("Tesseract OCR не установлен или недоступен")
        
        try:
            image = Image.open(file_path)
            logger.info(f"Изображение загружено: {image.size} пикселей")
            
            # Пробуем разные настройки OCR для лучшего распознавания
            configs = [
                '--oem 3 --psm 6',  # Стандартные настройки
                '--oem 3 --psm 3',  # Автоматическое определение страницы
                '--oem 1 --psm 6',  # Legacy OCR engine
                '--oem 3 --psm 8',  # Одна строка текста
                '--oem 3 --psm 13', # Необработанная строка
            ]
            
            best_text = ""
            for config in configs:
                try:
                    text = pytesseract.image_to_string(image, lang="rus+eng", config=config)
                    if text and len(text.strip()) > len(best_text.strip()):
                        best_text = text
                        logger.info(f"Успешное распознавание с конфигурацией: {config}")
                except Exception as e:
                    logger.warning(f"Ошибка OCR с конфигурацией {config}: {e}")
            
            return {
                'text': best_text,
                'pages': 1,
                'tables': 0,
                'images': 1
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки изображения: {e}")
            raise


# Глобальный экземпляр экстрактора
text_extractor = TextExtractor()


def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Удобная функция для извлечения текста из файла.
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        Словарь с результатами извлечения
    """
    return text_extractor.extract_text(file_path)


def get_supported_formats() -> List[str]:
    """Возвращает список поддерживаемых форматов файлов."""
    return list(text_extractor.supported_formats.keys())


def is_format_supported(file_path: str) -> bool:
    """Проверяет, поддерживается ли формат файла."""
    file_ext = os.path.splitext(file_path)[1].lower()
    return file_ext in text_extractor.supported_formats
