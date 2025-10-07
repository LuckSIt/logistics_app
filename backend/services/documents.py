import os
from typing import List, Any, Optional, Dict
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.graphics.shapes import Drawing, Rect, Line
from reportlab.graphics import renderPDF
from datetime import datetime
from pypdf import PdfReader, PdfWriter, PageObject
from io import BytesIO
# Workaround for ReportLab md5 compatibility on some Python/OpenSSL builds
try:
    import hashlib as _hashlib
    from reportlab.pdfbase import pdfdoc as _pdfdoc
    def _safe_md5(data=b"", **kwargs):
        h = _hashlib.md5()
        if data:
            try:
                h.update(data)
            except Exception:
                pass
        return h
    _pdfdoc.md5 = _safe_md5  # ignore unsupported usedforsecurity kwarg
except Exception:
    pass
from schemas import CalculateOption


def _get_value(opt: Any, key: str, default=None):
    try:
        if isinstance(opt, dict):
            return opt.get(key, default)
        return getattr(opt, key, default)
    except Exception:
        return default


def _normalize_option(opt: Any) -> dict:
    return {
        "supplier_name": _get_value(opt, "supplier_name") or _get_value(opt, "contractor_name", "-"),
        "price_rub": _get_value(opt, "price_rub") or _get_value(opt, "base_price_rub", 0.0),
        "price_usd": _get_value(opt, "price_usd"),
        "markup_percent": _get_value(opt, "markup_percent", 0.0),
        "discount_percent": _get_value(opt, "discount_percent") or _get_value(opt, "user_discount_percent", 0.0),
        "final_price_rub": _get_value(opt, "final_price_rub") or _get_value(opt, "total_price_rub") or _get_value(opt, "price_rub", 0.0),
        "border_point": _get_value(opt, "border_point"),
        "svh_name": _get_value(opt, "svh_name"),
        "arrival_station": _get_value(opt, "arrival_station"),
        "transit_time_days": _get_value(opt, "transit_time_days"),
        "validity_date": _get_value(opt, "validity_date"),
        # Дополнительные поля для детализации
        "rail_tariff_rub": _get_value(opt, "rail_tariff_rub"),
        "cbx_cost": _get_value(opt, "cbx_cost"),
        "auto_pickup_cost": _get_value(opt, "auto_pickup_cost"),
        "terminal_handling_cost": _get_value(opt, "terminal_handling_cost"),
        "security_cost": _get_value(opt, "security_cost"),
        "precarriage_cost": _get_value(opt, "precarriage_cost"),
    }


def generate_docx(options: List[CalculateOption], output_path: str, request_data: Optional[Dict[str, Any]] = None) -> str:
    doc = Document()
    doc.add_heading('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', 0)
    
    if request_data:
        _add_request_summary_docx(doc, request_data)
        doc.add_paragraph("Предлагаем рассмотреть:")
    
    # Генерация КП для каждого варианта
    for raw in options or []:
        opt = _normalize_option(raw)
        paragraphs = _build_commercial_proposal_paragraphs(request_data or {}, opt)
        for p in paragraphs:
            doc.add_paragraph(p)
        doc.add_paragraph("*******")  # Разделитель между вариантами
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_pdf(options: List[CalculateOption], output_path: str, request_data: Optional[Dict[str, Any]] = None) -> str:
    """Генерация профессионального PDF коммерческого предложения с использованием фирменного бланка"""
    # Создаем директорию только если путь содержит директорию
    dir_path = os.path.dirname(output_path)
    if dir_path:
        os.makedirs(dir_path, exist_ok=True)
    
    # Путь к фирменному бланку (в контейнере)
    blank_path = os.path.join(os.path.dirname(__file__), "..", "blank.pdf")
    
    # Проверяем наличие бланка
    if os.path.exists(blank_path):
        return _generate_pdf_with_letterhead(options, output_path, request_data, blank_path)
    else:
        # Если бланка нет, используем стандартную генерацию
        return _generate_pdf_standard(options, output_path, request_data)


def _generate_pdf_with_letterhead(options: List[CalculateOption], output_path: str, request_data: Optional[Dict[str, Any]] = None, blank_path: str = None) -> str:
    """Генерация PDF с использованием фирменного бланка"""
    # Регистрируем шрифт с поддержкой кириллицы
    cyrillic_font = _ensure_cyrillic_font()
    
    # Создаем контент в памяти
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        rightMargin=1.5*cm, 
        leftMargin=1.5*cm, 
        topMargin=4*cm,  # Чуть больше места сверху для бланка
        bottomMargin=2*cm
    )
    story = []
    
    # Получаем стили
    styles = getSampleStyleSheet()
    
    # Создаем улучшенные стили с поддержкой кириллицы
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.darkblue,
        fontName=cyrillic_font,
        leading=24
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=15,
        alignment=TA_LEFT,
        textColor=colors.darkblue,
        fontName=cyrillic_font,
        leading=18
    )
    
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=15,
        textColor=colors.darkblue,
        fontName=cyrillic_font,
        leading=16
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=4,
        alignment=TA_JUSTIFY,
        fontName=cyrillic_font,
        leading=14
    )
    
    price_style = ParagraphStyle(
        'PriceStyle',
        parent=normal_style,
        fontSize=12,
        textColor=colors.darkgreen,
        fontName=cyrillic_font,
        leading=16
    )
    
    # Добавляем логотип и заголовок
    story.append(_create_header_section(cyrillic_font))
    story.append(Spacer(1, 16))
    
    # Заголовок документа
    story.append(Paragraph("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", title_style))
    story.append(Spacer(1, 15))
    
    # Информация о компании в улучшенном формате
    story.append(_create_company_info_section(cyrillic_font, normal_style))
    story.append(Spacer(1, 20))
    
    # Дата и номер предложения
    current_date = datetime.now().strftime("%d.%m.%Y")
    story.append(Paragraph(f"<b>Дата:</b> {current_date}", normal_style))
    story.append(Paragraph(f"<b>Номер предложения:</b> КП-{datetime.now().strftime('%Y%m%d%H%M')}", normal_style))
    story.append(Spacer(1, 20))
    
    # Информация о запросе
    if request_data:
        request_sections = _create_request_info_section(request_data, cyrillic_font, header_style, normal_style)
        for section in request_sections:
            story.append(section)
        story.append(Spacer(1, 20))
    
    # Предложения
    story.append(Paragraph("ПРЕДЛАГАЕМЫЕ ВАРИАНТЫ:", header_style))
    story.append(Spacer(1, 10))
    
    # Генерация КП для каждого варианта
    for i, raw in enumerate(options or [], 1):
        opt = _normalize_option(raw)
        
        # Создаем блок для варианта
        variant_content = []
        
        # Заголовок варианта
        variant_content.append(Paragraph(f"Вариант {i}: {opt['supplier_name']}", subtitle_style))
        
        # Цена
        price_text = _format_price(opt)
        if price_text:
            variant_content.append(Paragraph(f"<b>Стоимость: {price_text}</b>", price_style))
        
        # Детализация предложения
        paragraphs = _build_commercial_proposal_paragraphs(request_data or {}, opt)
        for p in paragraphs:
            variant_content.append(Paragraph(f"• {p}", normal_style))
        
        # Добавляем содержимое варианта
        for content in variant_content:
            story.append(content)
        story.append(Spacer(1, 20))
    
    # Условия
    conditions_sections = _create_conditions_section(cyrillic_font, header_style, normal_style)
    for section in conditions_sections:
        story.append(section)
    story.append(Spacer(1, 20))
    
    # Контактная информация
    contact_sections = _create_contact_section(cyrillic_font, header_style, normal_style)
    for section in contact_sections:
        story.append(section)
    
    # Строим документ
    doc.build(story)
    
    # Теперь накладываем контент на фирменный бланк
    return _merge_with_letterhead(buffer, blank_path, output_path)


def _generate_pdf_standard(options: List[CalculateOption], output_path: str, request_data: Optional[Dict[str, Any]] = None) -> str:
    """Стандартная генерация PDF без фирменного бланка"""
    # Регистрируем шрифт с поддержкой кириллицы
    cyrillic_font = _ensure_cyrillic_font()
    
    # Создаем документ с улучшенными отступами
    doc = SimpleDocTemplate(
        output_path, 
        pagesize=A4, 
        rightMargin=1.5*cm, 
        leftMargin=1.5*cm, 
        topMargin=2*cm, 
        bottomMargin=2*cm
    )
    story = []
    
    # Получаем стили
    styles = getSampleStyleSheet()
    
    # Создаем улучшенные стили с поддержкой кириллицы
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.darkblue,
        fontName=cyrillic_font,
        leading=24
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=15,
        alignment=TA_LEFT,
        textColor=colors.darkblue,
        fontName=cyrillic_font,
        leading=18
    )
    
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=10,
        alignment=TA_LEFT,
        textColor=colors.black,
        fontName=cyrillic_font,
        leading=16
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
        fontName=cyrillic_font,
        leading=14
    )
    
    price_style = ParagraphStyle(
        'PriceStyle',
        parent=normal_style,
        fontSize=12,
        textColor=colors.darkgreen,
        fontName=cyrillic_font,
        leading=16
    )
    
    # Добавляем логотип и заголовок
    story.append(_create_header_section(cyrillic_font))
    story.append(Spacer(1, 16))
    
    # Заголовок документа
    story.append(Paragraph("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", title_style))
    story.append(Spacer(1, 15))
    
    # Информация о компании в улучшенном формате
    story.append(_create_company_info_section(cyrillic_font, normal_style))
    story.append(Spacer(1, 20))
    
    # Дата и номер предложения
    current_date = datetime.now().strftime("%d.%m.%Y")
    story.append(Paragraph(f"<b>Дата:</b> {current_date}", normal_style))
    story.append(Paragraph(f"<b>Номер предложения:</b> КП-{datetime.now().strftime('%Y%m%d%H%M')}", normal_style))
    story.append(Spacer(1, 20))
    
    # Информация о запросе
    if request_data:
        request_sections = _create_request_info_section(request_data, cyrillic_font, header_style, normal_style)
        for section in request_sections:
            story.append(section)
        story.append(Spacer(1, 20))
    
    # Предложения
    story.append(Paragraph("ПРЕДЛАГАЕМЫЕ ВАРИАНТЫ:", header_style))
    story.append(Spacer(1, 10))
    
    # Генерация КП для каждого варианта
    for i, raw in enumerate(options or [], 1):
        opt = _normalize_option(raw)
        
        # Создаем блок для варианта
        variant_content = []
        
        # Заголовок варианта
        variant_title = f"ВАРИАНТ {i}"
        if opt.get('supplier_name'):
            variant_title += f" - {opt['supplier_name']}"
        variant_content.append(Paragraph(variant_title, subtitle_style))
        
        # Генерируем текст КП
        if request_data:
            kp_paragraphs = _build_commercial_proposal_paragraphs(request_data, opt)
            for paragraph_text in kp_paragraphs:
                if paragraph_text.strip():
                    variant_content.append(Paragraph(paragraph_text, normal_style))
        
        # Цена
        price_text = _format_price(opt)
        if price_text:
            variant_content.append(Paragraph(f"<b>Стоимость:</b> {price_text}", price_style))
        
        # Добавляем содержимое варианта
        for content in variant_content:
            story.append(content)
        story.append(Spacer(1, 20))
    
    # Условия
    conditions_sections = _create_conditions_section(cyrillic_font, header_style, normal_style)
    for section in conditions_sections:
        story.append(section)
    story.append(Spacer(1, 20))
    
    # Контактная информация
    contact_sections = _create_contact_section(cyrillic_font, header_style, normal_style)
    for section in contact_sections:
        story.append(section)
    
    # Строим документ
    doc.build(story)
    return output_path


def _merge_with_letterhead(content_buffer: BytesIO, blank_path: str, output_path: str) -> str:
    """Объединяет контент с фирменным бланком"""
    try:
        # Читаем фирменный бланк
        blank_reader = PdfReader(blank_path)
        # ВАЖНО: перемотать буфер на начало, чтобы читать весь контент
        try:
            content_buffer.seek(0)
        except Exception:
            pass
        content_reader = PdfReader(content_buffer)
        
        # Создаем writer для объединения
        writer = PdfWriter()
        
        # Для каждой страницы контента
        for page_num in range(len(content_reader.pages)):
            content_page = content_reader.pages[page_num]

            if len(blank_reader.pages) > 0:
                # Берем соответствующую страницу бланка, если есть, иначе первую
                letter_page = blank_reader.pages[page_num] if page_num < len(blank_reader.pages) else blank_reader.pages[0]

                # Создаем новую пустую страницу по размеру бланка (чтобы не терять фон)
                new_page = PageObject.create_blank_page(
                    width=float(letter_page.mediabox.width),
                    height=float(letter_page.mediabox.height)
                )

                # Кладем бланк, затем контент
                new_page.merge_page(letter_page)
                new_page.merge_page(content_page)

                writer.add_page(new_page)
            else:
                # Если бланка нет, добавляем только контент
                writer.add_page(content_page)
        
        # Сохраняем результат
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return output_path
        
    except Exception as e:
        # Если что-то пошло не так, сохраняем только контент
        print(f"Ошибка при объединении с бланком: {e}")
        with open(output_path, 'wb') as output_file:
            output_file.write(content_buffer.getvalue())
    return output_path


def _format_price(opt: Dict[str, Any]) -> str:
    """Форматирование цены для отображения"""
    try:
        price_rub = opt.get('final_price_rub') or opt.get('price_rub')
        price_usd = opt.get('price_usd')
        
        if price_rub and price_usd:
            return f"{float(price_rub):,.0f} ₽ ({float(price_usd):,.0f} USD)"
        elif price_rub:
            return f"{float(price_rub):,.0f} ₽"
        elif price_usd:
            return f"{float(price_usd):,.0f} USD"
        else:
            return "по запросу"
    except Exception:
        return "по запросу"


# ----- helpers to build text -----
def _add_request_summary_docx(doc: Document, req: Dict[str, Any]):
    # Извлекаем данные из разных возможных мест в запросе
    origin_country = req.get('origin_country') or req.get('origin', {}).get('country') or ''
    origin_city = req.get('origin_city') or req.get('origin', {}).get('city') or ''
    destination_country = req.get('destination_country') or req.get('destination', {}).get('country') or ''
    destination_city = req.get('destination_city') or req.get('destination', {}).get('city') or ''
    basis = req.get('basis') or ''
    cargo_name = req.get('cargo_name') or ''
    weight_kg = req.get('weight_kg') or ''
    volume_m3 = req.get('volume_m3') or ''
    
    doc.add_paragraph(f"Адрес забора груза: {origin_country} {origin_city}")
    doc.add_paragraph(f"Базис поставки: {basis}")
    doc.add_paragraph(f"Адрес доставки груза: {destination_country} {destination_city}")
    if cargo_name:
        doc.add_paragraph(f"Наименование груза: {cargo_name}")
    if weight_kg:
        doc.add_paragraph(f"Вес груза: {weight_kg} кг")
    if volume_m3:
        doc.add_paragraph(f"Объем груза: {volume_m3} м³")


def _build_request_summary_lines(req: Dict[str, Any]) -> List[str]:
    # Извлекаем данные из разных возможных мест в запросе
    origin_country = req.get('origin_country') or req.get('origin', {}).get('country') or ''
    origin_city = req.get('origin_city') or req.get('origin', {}).get('city') or ''
    destination_country = req.get('destination_country') or req.get('destination', {}).get('country') or ''
    destination_city = req.get('destination_city') or req.get('destination', {}).get('city') or ''
    basis = req.get('basis') or ''
    cargo_name = req.get('cargo_name') or ''
    weight_kg = req.get('weight_kg') or ''
    volume_m3 = req.get('volume_m3') or ''
    
    lines = [
        f"Адрес забора груза: {origin_country} {origin_city}",
        f"Базис поставки: {basis}",
        f"Адрес доставки груза: {destination_country} {destination_city}",
    ]
    if cargo_name:
        lines.append(f"Наименование груза: {cargo_name}")
    if weight_kg:
        lines.append(f"Вес груза: {weight_kg} кг")
    if volume_m3:
        lines.append(f"Объем груза: {volume_m3} м³")
    return lines


def _build_commercial_proposal_paragraphs(req: Dict[str, Any], opt: Dict[str, Any]) -> List[str]:
    """Генерация коммерческого предложения согласно шаблонам из ТЗ"""
    import logging
    logger = logging.getLogger(__name__)
    
    try:
        transport_type = (req.get('transport_type') or opt.get('transport_type') or 'auto').lower()
        basis = (req.get('basis') or 'EXW').upper()

        logger.info(f"Генерируем КП для {transport_type} с базисом {basis}")
        logger.info(f"Данные запроса: {req}")
        logger.info(f"Данные тарифа: {opt}")
        
        # Используем новый менеджер шаблонов
        from kp_templates import kp_template_manager
        
        # Генерируем текст КП на основе шаблона
        kp_text = kp_template_manager.generate_kp_text(transport_type, basis, req, opt)
        logger.info(f"Сгенерированный текст КП: {kp_text[:200]}...")
        
        # Разбиваем на параграфы
        paragraphs = [p.strip() for p in kp_text.split('\n') if p.strip()]
        logger.info(f"Получено параграфов: {len(paragraphs)}")
        
        return paragraphs

    except Exception as e:
        logger.error(f"Ошибка генерации КП: {e}")
        return [f"Ошибка генерации КП: {str(e)}"]
    



def _draw_wrapped_line(c: canvas.Canvas, text: str, y: float, font: str, max_width: int = 500, leading: int = 14) -> float:
    if not text:
        return y
    words = text.split()
    line = ""
    for w in words:
        test = (line + " " + w).strip()
        if c.stringWidth(test, font, 10) > max_width:
            c.drawString(50, y, line)
            y -= leading
            line = w
        else:
            line = test
    if line:
        c.drawString(50, y, line)
        y -= leading
    return y


def _create_header_section(cyrillic_font: str) -> Table:
    """Создает секцию заголовка с логотипом и названием компании"""
    # Проверяем наличие логотипа
    logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "frontend", "public", "logo.png")
    
    if os.path.exists(logo_path):
        # Если логотип есть, создаем таблицу с логотипом
        from reportlab.platypus import Image
        try:
            logo = Image(logo_path, width=2*cm, height=2*cm)
            header_data = [
                [logo, "ООО «ВЕРЕС»", "Логистические услуги"],
                ["", "", "Коммерческое предложение"]
            ]
            
            header_table = Table(header_data, colWidths=[2*cm, 6*cm, 4*cm])
            header_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('ALIGN', (2, 0), (2, -1), 'RIGHT'),
                ('VALIGN', (0, 0), (0, -1), 'MIDDLE'),
                ('FONTNAME', (1, 0), (1, 0), cyrillic_font),
                ('FONTNAME', (2, 0), (2, -1), cyrillic_font),
                ('FONTSIZE', (1, 0), (1, 0), 16),
                ('FONTSIZE', (2, 0), (2, -1), 12),
                ('TEXTCOLOR', (1, 0), (1, 0), colors.darkblue),
                ('TEXTCOLOR', (2, 0), (2, -1), colors.grey),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
            ]))
            
            return header_table
        except Exception as e:
            print(f"Ошибка загрузки логотипа: {e}")
    
    # Если логотипа нет или ошибка, создаем простую таблицу
    header_data = [
        ["ООО «ВЕРЕС»", "Логистические услуги"],
        ["", "Коммерческое предложение"]
    ]
    
    header_table = Table(header_data, colWidths=[8*cm, 4*cm])
    header_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (0, 0), cyrillic_font),
        ('FONTNAME', (1, 0), (1, -1), cyrillic_font),
        ('FONTSIZE', (0, 0), (0, 0), 16),
        ('FONTSIZE', (1, 0), (1, -1), 12),
        ('TEXTCOLOR', (0, 0), (0, 0), colors.darkblue),
        ('TEXTCOLOR', (1, 0), (1, -1), colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
    ]))
    
    return header_table


def _create_company_info_section(cyrillic_font: str, normal_style: ParagraphStyle) -> Table:
    """Создает секцию с информацией о компании"""
    company_info = [
        [Paragraph("ООО «ВЕРЕС»", normal_style), Paragraph("", normal_style)],
        [Paragraph("Адрес:", normal_style), Paragraph("г. Москва, ул. Тверская, д. 15, стр. 1", normal_style)],
        [Paragraph("Телефон:", normal_style), Paragraph("+7 (495) 123-45-67", normal_style)],
        [Paragraph("Email:", normal_style), Paragraph("info@veres-logistics.ru", normal_style)],
        [Paragraph("Сайт:", normal_style), Paragraph("www.veres-logistics.ru", normal_style)],
        [Paragraph("ИНН:", normal_style), Paragraph("7701234567", normal_style)],
        [Paragraph("КПП:", normal_style), Paragraph("770101001", normal_style)],
        [Paragraph("ОГРН:", normal_style), Paragraph("1234567890123", normal_style)],
    ]
    
    company_table = Table(company_info, colWidths=[4*cm, 12*cm])
    company_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (0, -1), cyrillic_font),
        ('FONTNAME', (1, 0), (1, -1), cyrillic_font),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('BACKGROUND', (0, 0), (0, -1), colors.whitesmoke),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.grey),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.lightgrey),
    ]))
    
    return company_table


def _create_request_info_section(request_data: Dict[str, Any], cyrillic_font: str, header_style: ParagraphStyle, normal_style: ParagraphStyle) -> list:
    """Создает секцию с информацией о запросе"""
    sections = []
    sections.append(Paragraph("ПАРАМЕТРЫ ЗАПРОСА:", header_style))
    
    # Извлекаем данные из запроса
    origin_country = request_data.get('origin_country', '')
    origin_city = request_data.get('origin_city', '')
    destination_country = request_data.get('destination_country', '')
    destination_city = request_data.get('destination_city', '')
    basis = request_data.get('basis', '')
    cargo_name = request_data.get('cargo_name', '')
    weight_kg = request_data.get('weight_kg', '')
    volume_m3 = request_data.get('volume_m3', '')
    transport_type = request_data.get('transport_type', '')
    vehicle_type = request_data.get('vehicle_type', '')
    
    request_info = []
    if origin_city or origin_country:
        request_info.append(["Маршрут:", f"{origin_country} {origin_city} → {destination_country} {destination_city}"])
    if transport_type:
        transport_names = {
            'auto': 'Автомобильный',
            'rail': 'Железнодорожный', 
            'sea': 'Морской',
            'air': 'Авиа',
            'multimodal': 'Мультимодальный'
        }
        request_info.append(["Тип транспорта:", transport_names.get(transport_type, transport_type)])
    if basis:
        request_info.append(["Базис поставки:", basis])
    if cargo_name:
        request_info.append(["Наименование груза:", cargo_name])
    if weight_kg:
        request_info.append(["Вес:", f"{weight_kg} кг"])
    if volume_m3:
        request_info.append(["Объем:", f"{volume_m3} м³"])
    if vehicle_type:
        request_info.append(["Тип ТС:", vehicle_type])
    
    if request_info:
        request_table = Table(request_info, colWidths=[4*cm, 12*cm])
        request_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (0, -1), cyrillic_font),
            ('FONTNAME', (1, 0), (1, -1), cyrillic_font),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BACKGROUND', (0, 0), (-1, -1), colors.white),
            ('BACKGROUND', (0, 0), (0, -1), colors.whitesmoke),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.grey),
            ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ]))
        sections.append(request_table)
    
    return sections


def _create_conditions_section(cyrillic_font: str, header_style: ParagraphStyle, normal_style: ParagraphStyle) -> list:
    """Создает секцию с общими условиями"""
    sections = []
    sections.append(Paragraph("ОБЩИЕ УСЛОВИЯ:", header_style))
    
    conditions = [
        "• Срок действия предложения: 30 дней с даты выдачи",
        "• Валюта расчетов: Российский рубль (RUB) или Доллар США (USD)",
        "• Условия оплаты: 100% предоплата или по согласованию",
        "• Срок поставки: согласно указанному в предложении",
        "• Условия доставки: согласно выбранному базису поставки",
        "• Страхование: по согласованию с клиентом",
        "• Таможенное оформление: включено в стоимость (где применимо)",
        "• Ответственность: согласно международным конвенциям",
        "• Форс-мажор: согласно международной торговой практике"
    ]
    
    for condition in conditions:
        sections.append(Paragraph(condition, normal_style))
    
    return sections


def _create_contact_section(cyrillic_font: str, header_style: ParagraphStyle, normal_style: ParagraphStyle) -> list:
    """Создает секцию с контактной информацией"""
    sections = []
    sections.append(Paragraph("КОНТАКТНАЯ ИНФОРМАЦИЯ:", header_style))
    
    # Используем Paragraph, чтобы длинные подписи корректно переносились
    contact_info = [
        [Paragraph("Менеджер по работе с клиентами:", normal_style), Paragraph("Петров Петр Петрович", normal_style)],
        [Paragraph("Телефон:", normal_style), Paragraph("+7 (495) 123-45-67", normal_style)],
        [Paragraph("Мобильный:", normal_style), Paragraph("+7 (916) 123-45-67", normal_style)],
        [Paragraph("Email:", normal_style), Paragraph("petrov@veres-logistics.ru", normal_style)],
        [Paragraph("Время работы:", normal_style), Paragraph("Пн-Пт: 9:00-18:00 (МСК)", normal_style)],
    ]
    
    # Увеличиваем ширину левого столбца для длинных подписей
    contact_table = Table(contact_info, colWidths=[6*cm, 10*cm])
    contact_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (0, -1), cyrillic_font),
        ('FONTNAME', (1, 0), (1, -1), cyrillic_font),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('BACKGROUND', (0, 0), (0, -1), colors.whitesmoke),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.grey),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.lightgrey),
    ]))
    sections.append(contact_table)
    
    return sections


def _create_decorative_line() -> Drawing:
    """Создает декоративную линию для разделения секций"""
    drawing = Drawing(100, 20)
    drawing.add(Line(0, 10, 100, 10, strokeColor=colors.darkblue, strokeWidth=2))
    return drawing


def _ensure_cyrillic_font() -> str:
    """Register and return a font that supports Cyrillic. Falls back to Helvetica."""
    try:
        # If already registered
        if "DejaVuSans" in pdfmetrics.getRegisteredFontNames():
            return "DejaVuSans"
        
        # Try common font locations - prioritize fonts that support Cyrillic
        candidates = [
            os.getenv("FONT_PATH"),
            r"C:\Windows\Fonts\calibri.ttf",  # Calibri has good Cyrillic support
            r"C:\Windows\Fonts\verdana.ttf",  # Verdana has good Cyrillic support
            r"C:\Windows\Fonts\tahoma.ttf",   # Tahoma has good Cyrillic support
            r"C:\Windows\Fonts\arial.ttf",    # Arial has limited Cyrillic support
            r"C:\Windows\Fonts\times.ttf",    # Times New Roman
            "DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/local/share/fonts/DejaVuSans.ttf",
        ]
        
        for path in candidates:
            if path and os.path.exists(path):
                try:
                    name = os.path.splitext(os.path.basename(path))[0]
                    pdfmetrics.registerFont(TTFont(name, path))
                    print(f"✅ Используется шрифт: {name} из {path}")
                    return name
                except Exception as e:
                    print(f"❌ Ошибка загрузки шрифта {path}: {e}")
                    continue
                    
        # Если не удалось загрузить TTF, используем встроенный шрифт
        print("⚠️ Используется встроенный шрифт Helvetica")
        return "Helvetica"
        
    except Exception as e:
        print(f"❌ Ошибка при работе с шрифтами: {e}")
        return "Helvetica"


