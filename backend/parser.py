import os
import re
import json
import logging
from typing import List, Dict, Any, Optional
from datetime import datetime, date
import fitz
import docx
import openpyxl
import ollama
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)

# Настройка Tesseract OCR
try:
    # Попробуем несколько возможных путей к Tesseract
    possible_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract"
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            logger.info(f"Tesseract found at: {path}")
            break
    else:
        logger.warning("Tesseract not found. OCR for images will be unavailable.")
except Exception as e:
    logger.error(f"Error setting up Tesseract: {e}")


class LLMParser:
    def __init__(self, model: str = "mistral"):
        self.model = model
    
    def extract_text_from_file(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == ".pdf":
                return self._parse_pdf(file_path)
            elif ext == ".docx":
                return self._parse_docx(file_path)
            elif ext in [".xlsx", ".xls"]:
                return self._parse_xlsx(file_path)
            elif ext in [".png", ".jpg", ".jpeg"]:
                return self._parse_image(file_path)
            else:
                logger.warning(f"Unsupported file format: {ext}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _parse_pdf(self, path: str) -> str:
        text = ""
        try:
            with fitz.open(path) as doc:
                for page in doc:
                    text += page.get_text()
        except Exception as e:
            logger.error(f"Error parsing PDF {path}: {e}")
        return text

    def _parse_docx(self, path: str) -> str:
        try:
            d = docx.Document(path)
            text_parts = []
            
            # Извлекаем текст из параграфов
            for p in d.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text.strip())
            
            # Извлекаем текст из таблиц
            for table in d.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing DOCX {path}: {e}")
            return ""
    
    def _parse_xlsx(self, path: str) -> str:
        try:
            wb = openpyxl.load_workbook(path)
            text = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
            for row in ws.iter_rows(values_only=True):
                text += " ".join([str(c) for c in row if c]) + "\n"
            return text
        except Exception as e:
            logger.error(f"Error parsing Excel {path}: {e}")
            return ""
    
    def _parse_image(self, path: str) -> str:
        """Extract text from images using OCR"""
        try:
            logger.info(f"Starting OCR for image: {path}")
            
            # Check if Tesseract is available
            try:
                pytesseract.get_tesseract_version()
            except Exception as e:
                logger.error(f"Tesseract is not available: {e}")
                return ""
            
            # Open and process the image
            image = Image.open(path)
            logger.info(f"Image loaded: {image.size} pixels")
            
            # Try different OCR configurations for better recognition
            configs = [
                '--oem 3 --psm 6',  # Standard settings
                '--oem 3 --psm 3',  # Automatic page segmentation
                '--oem 1 --psm 6',  # Legacy OCR engine
                '--oem 3 --psm 8',  # Single text line
                '--oem 3 --psm 13', # Raw line
            ]
            
            best_text = ""
            for config in configs:
                try:
                    text = pytesseract.image_to_string(image, lang="rus+eng", config=config)
                    if text and len(text.strip()) > len(best_text.strip()):
                        best_text = text
                        logger.info(f"Successful recognition with config: {config}")
                except Exception as e:
                    logger.warning(f"OCR error with config {config}: {e}")
            
            if best_text:
                logger.info(f"OCR extracted {len(best_text)} characters from image")
                return best_text
            else:
                logger.warning("Could not extract text from image")
                return ""
                
        except Exception as e:
            logger.error(f"Error processing image {path}: {e}")
        return ""

    def extract_tariff_data(self, text: str, transport_type: str = "auto") -> List[Dict[str, Any]]:
        """
        Извлекает данные тарифов из текста через LLM в JSON, затем конвертирует в формат приложения
        """
        if not text.strip():
            return []
        
        # Шаг 1: Извлекаем сырые данные в JSON
        raw_json_data = self._extract_raw_json_data(text, transport_type)
        if not raw_json_data:
            logger.error("Не удалось извлечь JSON данные из LLM")
            return []
        
        # Шаг 2: Конвертируем JSON в формат приложения
        formatted_data = self._convert_json_to_app_format(raw_json_data, transport_type)
        return formatted_data
    
    def _extract_raw_json_data(self, text: str, transport_type: str) -> List[Dict[str, Any]]:
        """Извлекает сырые JSON данные из текста через LLM"""
        text_sample = text[:4000] if len(text) > 4000 else text
        
        prompt = f"""
        Извлеки данные о транспортировке и тарифах из текста.
        Верни ТОЛЬКО JSON массив, где каждый объект = отдельный маршрут/тариф.
        
        Формат ответа:
        [
          {{
            "transport_type": "{transport_type}",
            "basis": "EXW",
            "origin_country": "string|null",
            "origin_city": "string|null", 
            "destination_country": "string|null",
            "destination_city": "string|null",
            "vehicle_type": "string|null",
            "price_rub": number|null,
            "price_usd": number|null,
            "validity_date": "string|null",
            "transit_time_days": "string|null",
            "conditions": "string|null"
          }}
        ]

        Текст для анализа:
        \"\"\"{text_sample}\"\"\"
        """
        
        try:
            resp = ollama.chat(
                model=self.model, 
                messages=[{"role": "user", "content": prompt}],
                options={"temperature": 0.1, "num_predict": 3000}
            )
            content = resp["message"]["content"]
            # Сохраняем ответ для отладки
            with open("llm_response_debug.txt", "w", encoding="utf-8") as f:
                f.write(content)
            logger.info("LLM response saved to llm_response_debug.txt for debugging")
            
            json_data = self._extract_json_from_response(content)
            return json_data if json_data else []
        except Exception as e:
            logger.error(f"Ошибка в LLM обработке: {e}")
            return []
    
    def _convert_json_to_app_format(self, raw_data: List[Dict[str, Any]], transport_type: str) -> List[Dict[str, Any]]:
        """Конвертирует сырые JSON данные в формат приложения"""
        formatted_data = []
        
        for item in raw_data:
            try:
                # Нормализуем данные
                formatted_item = self._normalize_llm_data(item, transport_type)
                if formatted_item:
                    formatted_data.append(formatted_item)
            except Exception as e:
                logger.warning(f"Ошибка нормализации элемента: {e}, элемент: {item}")
                continue
        
        logger.info(f"Успешно конвертировано {len(formatted_data)} записей из {len(raw_data)}")
        return formatted_data
    
    def _extract_json_from_response(self, content: str) -> Optional[List[Dict]]:
        """Извлекает JSON из ответа LLM с улучшенной обработкой ошибок"""
        candidates = []
        
        # Стратегия 1: Ищем JSON в блоках кода
        json_blocks = re.findall(r"```json\s*(.*?)\s*```", content, re.S)
        candidates.extend(json_blocks)
        
        # Стратегия 2: Ищем JSON массив
        array_matches = re.findall(r'\[.*?\]', content, re.S)
        candidates.extend(array_matches)
        
        # Стратегия 3: Ищем JSON объект
        obj_matches = re.findall(r'\{.*?\}', content, re.S)
        candidates.extend(obj_matches)
        
        # Пробуем каждый кандидат
        for i, candidate in enumerate(candidates):
            try:
                # Очищаем кандидата
                candidate = candidate.strip()
                if not candidate:
                    continue
                
                # Исправляем ошибки JSON
                fixed_candidate = self._fix_json_errors(candidate)
                
                # Пытаемся распарсить
                parsed = json.loads(fixed_candidate)
                
                # Проверяем, что это массив
                if isinstance(parsed, list):
                    logger.info(f"Успешно извлечен JSON массив из кандидата {i+1}")
                    return parsed
                elif isinstance(parsed, dict):
                    # Если это объект, оборачиваем в массив
                    logger.info(f"Извлечен JSON объект из кандидата {i+1}, оборачиваем в массив")
                    return [parsed]
                    
            except Exception as e:
                logger.debug(f"Кандидат {i+1} не прошел парсинг: {e}")
                continue
        
        # Если ничего не сработало, сохраняем для отладки
        logger.error("Не удалось извлечь валидный JSON из ответа LLM")
        with open("llm_response_debug.txt", "w", encoding="utf-8") as f:
            f.write(content)
        logger.info("LLM response saved to llm_response_debug.txt for debugging")
        
        return None
    
    def _fix_json_errors(self, json_str: str) -> str:
        """Исправляет распространенные ошибки в JSON от LLM"""
        # Убираем управляющие символы
        json_str = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_str)
        
        # Убираем лишние запятые перед закрывающими скобками
        json_str = re.sub(r',\s*}', '}', json_str)
        json_str = re.sub(r',\s*]', ']', json_str)
        
        # Заменяем одинарные кавычки на двойные
        json_str = re.sub(r"'([^']*)'", r'"\1"', json_str)
        
        # Исправляем значения без кавычек (кроме чисел, null, true, false)
        json_str = re.sub(r':\s*([^",}\]]+?)([,}\]])', r': "\1"\2', json_str)
        
        # Исправляем числа с дефисами (диапазоны) - оставляем как строки
        json_str = re.sub(r':\s*"(\d+-\d+)"', r': "\1"', json_str)
        
        # Исправляем специальные значения
        json_str = re.sub(r':\s*"(null|true|false)"', r': \1', json_str)
        json_str = re.sub(r':\s*"(null|true|false)"', r': \1', json_str)
        
        # Исправляем числа
        json_str = re.sub(r':\s*"(\d+\.?\d*)"', r': \1', json_str)
        
        # Убираем лишние пробелы и переносы строк
        json_str = json_str.strip()
        
        return json_str
    
    def _normalize_llm_data(self, data: Dict[str, Any], transport_type: str) -> Optional[Dict[str, Any]]:
        try:
            price_rub = self._parse_number(data.get("price_rub"))
            price_usd = self._parse_number(data.get("price_usd"))
            transit_time_days = self._parse_number(data.get("transit_time_days"))
            
            validity_date = self._parse_date(data.get("validity_date"))
            
            normalized = {
                "transport_type": transport_type,
                "basis": data.get("basis", "EXW"),
                "origin_country": self._clean_string(data.get("origin_country")),
                "origin_city": self._clean_string(data.get("origin_city")),
                "destination_country": self._clean_string(data.get("destination_country")),
                "destination_city": self._clean_string(data.get("destination_city")),
                "vehicle_type": self._clean_string(data.get("vehicle_type")),
                "price_rub": price_rub,
                "price_usd": price_usd,
                "validity_date": validity_date,
                "transit_time_days": transit_time_days,
            }
            
            normalized = {k: v for k, v in normalized.items() if v is not None and v != ""}
            
            if (normalized.get("origin_city") or normalized.get("destination_city") or 
                normalized.get("price_rub") or normalized.get("price_usd")):
                return normalized
            
        except Exception as e:
            logger.error(f"Error normalizing LLM data: {e}")
        
        return None
    
    def _parse_number(self, value: Any) -> Optional[float]:
        if value is None:
            return None
        
        try:
            if isinstance(value, (int, float)):
                return float(value)
            
            str_value = str(value).strip()
            
            # Обрабатываем диапазоны типа "20-25"
            if '-' in str_value and str_value.count('-') == 1:
                parts = str_value.split('-')
                if len(parts) == 2:
                    try:
                        start = float(re.sub(r'[^\d.,]', '', parts[0]).replace(',', '.'))
                        end = float(re.sub(r'[^\d.,]', '', parts[1]).replace(',', '.'))
                        return (start + end) / 2  # Возвращаем среднее значение
                    except (ValueError, TypeError):
                        pass
            
            str_value = re.sub(r'[^\d.,\-]', '', str_value)
            str_value = str_value.replace(',', '.')
            
            if str_value:
                return float(str_value)
        except (ValueError, TypeError):
            pass
        
        return None
    
    def _parse_date(self, value: Any) -> Optional[str]:
        if value is None:
            return None
        
        try:
            if isinstance(value, str):
                value = value.strip()
                if not value:
                    return None
                
                date_formats = [
                    "%d.%m.%Y", "%d/%m/%Y", "%Y-%m-%d", 
                    "%d.%m.%y", "%d/%m/%y", "%y-%m-%d"
                ]
                
                for fmt in date_formats:
                    try:
                        parsed_date = datetime.strptime(value, fmt)
                        return parsed_date.strftime("%Y-%m-%d")
                    except ValueError:
                        continue
                
                return value
        except Exception:
            pass
        
        return None
    
    def _clean_string(self, value: Any) -> Optional[str]:
        if value is None:
            return None
        
        cleaned = str(value).strip()
        return cleaned if cleaned else None
    
    def parse_file(self, file_path: str, transport_type: str = "auto") -> List[Dict[str, Any]]:
        logger.info(f"Starting LLM parsing of file: {file_path}")
        
        text = self.extract_text_from_file(file_path)
        if not text.strip():
            logger.warning(f"Could not extract text from file: {file_path}")
            return []
        
        return self.extract_tariff_data(text, transport_type)


# Compatibility functions
def parse_pdf(path):
    parser = LLMParser()
    return parser._parse_pdf(path)


def parse_docx(path):
    parser = LLMParser()
    return parser._parse_docx(path)


def parse_xlsx(path):
    parser = LLMParser()
    return parser._parse_xlsx(path)


def parse_file(path):
    parser = LLMParser()
    return parser.extract_text_from_file(path)


def parse_image(path):
    """Compatibility function for image parsing"""
    parser = LLMParser()
    return parser._parse_image(path)


def extract_info(text, model="mistral"):
    parser = LLMParser(model)
    return parser.extract_tariff_data(text)


def process_folder(folder, out_file="results.json"):
    parser = LLMParser()
    results = []
    
    for f in os.listdir(folder):
        if f.lower().endswith((".pdf", ".docx", ".xlsx", ".xls", ".png", ".jpg", ".jpeg")):
            path = os.path.join(folder, f)
            print(f"Processing {path} ...")
            
            text = parser.extract_text_from_file(path)
            if text:
                data = parser.extract_tariff_data(text)
            results.append({
                "file": f,
                "routes": data
            })

    with open(out_file, "w", encoding="utf-8") as fout:
        json.dump(results, fout, ensure_ascii=False, indent=2)

    print(f"\nDone! Results saved to {out_file}")


if __name__ == "__main__":
    folder = "files"
    process_folder(folder, out_file="results.json")
